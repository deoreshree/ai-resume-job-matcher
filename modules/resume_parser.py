"""Resume parser: PDF/DOCX/TXT extraction plus hybrid NLP/heuristic structuring."""

from __future__ import annotations

import io
import logging
import re
from typing import Any, Optional

from pypdf import PdfReader
from pypdf.errors import PdfReadError

from config import SPACY_MODEL
from modules.skill_extractor import extract_skills
from utils.helpers import (
    estimate_years_of_experience,
    find_cgpa_or_percentage,
    find_degree,
    find_field_of_study,
    find_location,
    parse_dates_from_block,
)
from utils.text_cleaner import clean_text
from utils.validators import ResumeValidationError, validate_resume_file

logger = logging.getLogger(__name__)

EMAIL_RE = re.compile(r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b")
PHONE_RE = re.compile(
    r"(?:(?:\+?\d{1,3}[\s.-]*)?(?:\(?\d{2,4}\)?[\s.-]*)?\d{3,4}[\s.-]?\d{4})"
)
URL_RE = re.compile(
    r"(https?://[^\s)]+|www\.[^\s)]+|linkedin\.com/in/[^\s)]+|github\.com/[^\s)]+)",
    re.IGNORECASE,
)

_HEADING_ALIASES: dict[str, tuple[str, ...]] = {
    "summary": (
        "summary",
        "profile",
        "objective",
        "about me",
        "professional summary",
        "executive summary",
        "career objective",
        "career summary",
        "overview",
    ),
    "experience": (
        "experience",
        "work experience",
        "employment",
        "professional experience",
        "work history",
        "career history",
        "employment history",
        "professional background",
    ),
    "internships": ("internship", "internships", "internship experience"),
    "education": (
        "education",
        "academic background",
        "academics",
        "qualifications",
        "academic qualifications",
        "educational background",
        "academic credentials",
    ),
    "skills": (
        "skills",
        "technical skills",
        "core competencies",
        "technologies",
        "tech stack",
        "key skills",
        "competencies",
        "area of expertise",
        "technical competencies",
        "skills & competencies",
    ),
    "projects": (
        "projects",
        "personal projects",
        "academic projects",
        "key projects",
        "technical projects",
        "portfolio projects",
        "side projects",
    ),
    "certifications": (
        "certifications",
        "certificates",
        "licenses",
        "certifications & licenses",
    ),
    "achievements": (
        "achievements",
        "awards",
        "honors",
        "accomplishments",
        "awards & honors",
    ),
    "languages": (
        "languages",
        "known languages",
        "language proficiency",
        "languages spoken",
    ),
    "interests": ("interests", "hobbies", "hobbies & interests", "personal interests"),
    "extracurricular": (
        "extracurricular",
        "extracurricular activities",
        "volunteer",
        "volunteering",
        "community involvement",
    ),
    "contact": ("contact", "contact information", "personal information", "contact details"),
}

_HEADING_LOOKUP: dict[str, str] = {}
for _section, _aliases in _HEADING_ALIASES.items():
    for _alias in _aliases:
        _HEADING_LOOKUP[_alias] = _section

_NAME_STOP = re.compile(
    r"resume|curriculum|vitae|contact|email|phone|objective|summary|profile|education|experience|skills",
    re.IGNORECASE,
)
_STOPWORDS = {
    "the",
    "and",
    "for",
    "with",
    "from",
    "that",
    "this",
    "using",
    "used",
    "into",
    "over",
    "your",
    "their",
    "have",
    "has",
    "were",
    "been",
    "also",
    "etc",
}

_nlp = None
_nlp_load_attempted = False


def get_nlp():
    """Load spaCy once. Returns None if unavailable."""
    global _nlp, _nlp_load_attempted
    if _nlp_load_attempted:
        return _nlp
    _nlp_load_attempted = True
    try:
        import spacy

        _nlp = spacy.load(SPACY_MODEL)
    except Exception:
        logger.warning(
            "spaCy model '%s' is unavailable; using heuristic parsing only.",
            SPACY_MODEL,
        )
        _nlp = None
    return _nlp


def extract_text(file_bytes: bytes, filename: str) -> str:
    """Return cleaned raw text from a PDF, DOCX, or TXT resume."""
    validate_resume_file(filename, file_bytes)
    name = filename.lower()
    try:
        if name.endswith(".pdf"):
            raw = _extract_pdf(file_bytes)
        elif name.endswith(".docx"):
            raw = _extract_docx(file_bytes)
        elif name.endswith(".txt"):
            raw = _extract_txt(file_bytes)
        else:
            raise ResumeValidationError(
                "Unsupported file type. Please upload a PDF, DOCX, or TXT resume."
            )
    except ResumeValidationError:
        raise
    except PdfReadError as exc:
        raise ResumeValidationError(
            "The PDF could not be read. It may be damaged, encrypted, or image-only."
        ) from exc
    except Exception as exc:
        if filename.lower().endswith(".docx"):
            raise ResumeValidationError(
                "The DOCX file could not be read. Please export it again and retry."
            ) from exc
        raise ResumeValidationError(
            "The resume could not be read. Please try another file."
        ) from exc

    cleaned = clean_text(raw)
    
    # Scanned / image-based PDF check: if less than 20 alphanumeric chars, flag ocr_required
    alphanumeric = re.sub(r"\W", "", cleaned)
    if len(alphanumeric) < 20:
        raise ResumeValidationError(
            "This resume appears to be image-based. OCR is required to extract its text accurately.",
        )
    return cleaned


def parse_resume(file_bytes: bytes, filename: str) -> dict[str, Any]:
    """Parse a resume file into a structured dictionary."""
    try:
        text = extract_text(file_bytes, filename)
        return parse_resume_text(text)
    except ResumeValidationError as exc:
        if "image-based" in str(exc).lower() or "ocr" in str(exc).lower():
            return _empty_ocr_required_resume(str(exc))
        raise


def _empty_ocr_required_resume(message: str) -> dict[str, Any]:
    """Return structured schema for scanned/image-based resumes requiring OCR."""
    empty_candidate = {
        "name": None,
        "email": None,
        "phone": None,
        "location": None,
        "linkedin": None,
        "github": None,
        "portfolio": None,
    }
    empty_categorized_skills = {
        "programming_languages": [],
        "frameworks": [],
        "libraries": [],
        "databases": [],
        "cloud": [],
        "devops": [],
        "ai_ml": [],
        "tools": [],
        "soft_skills": [],
        "other": [],
    }
    return {
        "status": "ocr_required",
        "message": message,
        "candidate": empty_candidate,
        "contact": empty_candidate,
        "personal_information": empty_candidate,
        "name": None,
        "email": None,
        "phone": None,
        "location": None,
        "linkedin": None,
        "github": None,
        "portfolio": None,
        "urls": [],
        "summary": "",
        "skills": empty_categorized_skills,
        "skills_list": [],
        "skills_by_category": {},
        "education": [],
        "experience": [],
        "projects": [],
        "certifications": [],
        "achievements": [],
        "languages": [],
        "interests": [],
        "internships": [],
        "extracurricular": [],
        "keywords": [],
        "years_experience": 0.0,
        "sections": {},
        "raw_text": "",
        "parsing": {
            "status": "ocr_required",
            "confidence": 0.0,
            "warnings": [message],
        },
    }


def parse_resume_text(text: str) -> dict[str, Any]:
    """Parse cleaned resume text into the standardized structured schema."""
    cleaned = clean_text(text)
    if not cleaned:
        raise ResumeValidationError("The resume appears to be empty after text extraction.")

    sections = split_sections(cleaned)
    email = _first_match(EMAIL_RE, cleaned)
    phone = _normalize_phone(_first_match(PHONE_RE, cleaned))
    urls = _unique(URL_RE.findall(cleaned))
    name = extract_name(cleaned, sections)

    linkedin = next((url for url in urls if "linkedin.com" in url.casefold()), None)
    github = next((url for url in urls if "github.com" in url.casefold()), None)
    portfolio = next((url for url in urls if url not in {linkedin, github}), None)
    location = find_location(sections.get("preamble", "") or cleaned[:600])

    candidate = {
        "name": name,
        "email": email,
        "phone": phone,
        "location": location,
        "linkedin": linkedin,
        "github": github,
        "portfolio": portfolio,
    }

    summary = sections.get("summary", "").strip()

    education = parse_education(sections.get("education", ""))
    experience = parse_experience_entries(sections.get("experience", ""))
    internships = parse_experience_entries(sections.get("internships", ""))
    projects = parse_project_entries(sections.get("projects", ""))
    certifications = _bullet_or_line_items(sections.get("certifications", ""))
    achievements = _bullet_or_line_items(sections.get("achievements", ""))
    languages = _bullet_or_line_items(sections.get("languages", ""))
    interests = _bullet_or_line_items(sections.get("interests", ""))
    extracurricular = _bullet_or_line_items(sections.get("extracurricular", ""))

    section_skills = extract_skill_phrases(sections.get("skills", ""), cleaned)
    skill_details = extract_skills(cleaned)
    skills_list = _unique(skill_details["skills"] + section_skills)
    skills_categorized = skill_details.get("categorized_skills", {})
    skills_by_category_active = skill_details.get("by_category", {})

    keywords = extract_keywords(cleaned, skills_list)
    years = estimate_years_of_experience(sections.get("experience", "") or cleaned)

    # Validation & Parsing Confidence
    confidence = 0.95
    warnings: list[str] = []

    if not name:
        confidence -= 0.15
        warnings.append("Candidate name could not be confidently identified.")
    if not email:
        confidence -= 0.15
        warnings.append("Contact email could not be detected.")
    if not skills_list:
        confidence -= 0.15
        warnings.append("No technical skills could be confidently extracted.")
    if not education:
        confidence -= 0.10
        warnings.append("Education section could not be confidently identified.")
    if not experience and not internships:
        confidence -= 0.10
        warnings.append("No explicit work experience section detected.")

    confidence = round(max(0.20, min(0.99, confidence)), 2)
    parsing_status = "success" if confidence >= 0.70 else "partial"

    return {
        # Standardized schema fields
        "candidate": candidate,
        "summary": summary,
        "skills": skills_list,  # List of string names for assertion and matching compatibility
        "categorized_skills": skills_categorized,
        "education": education,
        "experience": experience,
        "projects": projects,
        "certifications": certifications,
        "achievements": achievements,
        "languages": languages,
        "interests": interests,
        "internships": internships,
        "extracurricular": extracurricular,
        "parsing": {
            "status": parsing_status,
            "confidence": confidence,
            "warnings": warnings,
        },
        # Backward-compatibility accessors
        "name": name,
        "email": email,
        "phone": phone,
        "location": location,
        "linkedin": linkedin,
        "github": github,
        "portfolio": portfolio,
        "urls": urls,
        "contact": candidate,
        "personal_information": candidate,
        "skills_list": skills_list,
        "skills_by_category": skills_by_category_active,
        "skill_details": skill_details,
        "keywords": keywords,
        "years_experience": years,
        "sections": sections,
        "raw_text": cleaned,
    }


def split_sections(text: str) -> dict[str, str]:
    """Split resume text into known sections using heading detection."""
    lines = text.split("\n")
    buckets: dict[str, list[str]] = {"preamble": []}
    current = "preamble"

    for line in lines:
        heading = _detect_heading(line)
        if heading:
            current = heading
            buckets.setdefault(current, [])
            continue
        buckets.setdefault(current, []).append(line)

    return {key: "\n".join(values).strip() for key, values in buckets.items() if "".join(values).strip()}


def extract_name(text: str, sections: Optional[dict[str, str]] = None) -> Optional[str]:
    """Guess the candidate name from the header and optional NER."""
    header_lines = [line.strip(" |-") for line in text.split("\n")[:12] if line.strip(" |-")]
    header = "\n".join(header_lines)
    nlp = get_nlp()
    if nlp is not None:
        doc = nlp(header[:1500])
        persons = [ent.text.strip() for ent in doc.ents if ent.label_ == "PERSON"]
        for person in persons:
            candidate = person.split("\n")[0].strip(" |-")
            if _looks_like_name(candidate):
                return candidate

    for candidate in header_lines:
        if _looks_like_name(candidate):
            return candidate
    return None


def extract_skill_phrases(skills_section: str, full_text: str) -> list[str]:
    """Collect skill-like phrases from the dedicated skills section."""
    source = skills_section or ""
    parts = re.split(r"[,|/•;]|\n| - ", source)
    skills: list[str] = []
    for part in parts:
        item = re.sub(r"^(?:technical|soft|skills?)\s*[:\-]\s*", "", part.strip(), flags=re.I)
        item = item.strip(" -:\t")
        if 1 < len(item) <= 40 and not EMAIL_RE.search(item) and not item.lower().startswith("http"):
            if not re.fullmatch(r"[\W_]+", item):
                skills.append(_title_skill(item))
    return _unique(skills)[:80]


def extract_keywords(text: str, skills: list[str]) -> list[str]:
    """Keywords from noun chunks when spaCy is available, else skills + notable tokens."""
    keywords = list(skills)
    nlp = get_nlp()
    if nlp is not None:
        doc = nlp(text[:8000])
        for chunk in doc.noun_chunks:
            phrase = clean_text(chunk.text)
            if _usable_keyword(phrase):
                keywords.append(_title_skill(phrase))
    else:
        for token in re.findall(r"\b[A-Z][A-Za-z0-9+#.]{2,}\b", text):
            if _usable_keyword(token):
                keywords.append(token)
    return _unique(keywords)[:50]


def parse_education(section_text: str) -> list[dict[str, Any]]:
    if not section_text:
        return []
    blocks = _split_blocks(section_text)
    entries: list[dict[str, Any]] = []
    for block in blocks:
        degree = find_degree(block)
        field = find_field_of_study(block)
        institution = _guess_institution(block)
        location = find_location(block)
        cgpa_or_pct = find_cgpa_or_percentage(block)
        start_date, end_date, duration = parse_dates_from_block(block)

        entries.append(
            {
                "degree": degree or "",
                "field": field or "",
                "institution": institution or "",
                "location": location or "",
                "start_date": start_date or "",
                "end_date": end_date or "",
                "percentage_or_cgpa": cgpa_or_pct or "",
                # Legacy compatibility
                "raw": block,
                "year": end_date or start_date or None,
            }
        )
    return entries


def parse_experience_entries(section_text: str) -> list[dict[str, Any]]:
    if not section_text:
        return []
    entries: list[dict[str, Any]] = []
    for block in _split_blocks(section_text):
        lines = [line for line in block.split("\n") if line.strip()]
        if not lines:
            continue
        headline = lines[0]
        bullets = [line.lstrip("- ").strip() for line in lines[1:] if line.strip()]
        
        job_title = headline
        company = _guess_organization(headline) or ""
        location = find_location(block) or ""
        start_date, end_date, duration = parse_dates_from_block(block)
        technologies = extract_skills(block)["skills"]

        entries.append(
            {
                "job_title": job_title,
                "company": company,
                "location": location,
                "start_date": start_date or "",
                "end_date": end_date or "",
                "description": bullets,
                "technologies": technologies,
                # Legacy compatibility
                "raw": block,
                "title": headline,
                "organization": company,
                "bullets": bullets,
                "responsibilities": bullets,
                "duration": duration,
            }
        )
    return entries


def parse_project_entries(section_text: str) -> list[dict[str, Any]]:
    if not section_text:
        return []
    entries: list[dict[str, Any]] = []
    for block in _split_blocks(section_text):
        lines = [line for line in block.split("\n") if line.strip()]
        if not lines:
            continue
        title = lines[0].lstrip("- ").strip()
        description = " ".join(line.lstrip("- ").strip() for line in lines[1:])
        bullets = [line.lstrip("- ").strip() for line in lines[1:] if line.strip()]
        technologies = extract_skills(block)["skills"]
        link = _first_match(URL_RE, block) or ""

        entries.append(
            {
                "name": title,
                "description": description,
                "technologies": technologies,
                "role": "Developer / Contributor",
                "link": link,
                # Legacy compatibility
                "title": title,
                "bullets": bullets,
                "raw": block,
            }
        )
    return entries


def _extract_pdf(file_bytes: bytes) -> str:
    reader = PdfReader(io.BytesIO(file_bytes))
    if getattr(reader, "is_encrypted", False):
        try:
            reader.decrypt("")
        except Exception as exc:
            raise ResumeValidationError(
                "This PDF is password-protected. Please upload an unlocked copy."
            ) from exc
    pages: list[str] = []
    for page in reader.pages:
        pages.append(page.extract_text() or "")
    return "\n".join(pages)


def _extract_docx(file_bytes: bytes) -> str:
    from docx import Document
    from docx.opc.exceptions import PackageNotFoundError

    try:
        document = Document(io.BytesIO(file_bytes))
    except PackageNotFoundError as exc:
        raise ResumeValidationError(
            "The DOCX file could not be read. Please upload a valid Word document."
        ) from exc

    parts: list[str] = []
    for paragraph in document.paragraphs:
        if paragraph.text:
            parts.append(paragraph.text)
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _extract_txt(file_bytes: bytes) -> str:
    """Extract text from a .txt file using UTF-8 with fallbacks."""
    for encoding in ("utf-8", "latin-1", "cp1252"):
        try:
            return file_bytes.decode(encoding)
        except UnicodeDecodeError:
            continue
    return file_bytes.decode("utf-8", errors="replace")


def _detect_heading(line: str) -> Optional[str]:
    stripped = line.strip().strip(":").strip()
    if not stripped or len(stripped) > 48:
        return None
    if stripped.startswith("-"):
        return None
    key = re.sub(r"[^a-zA-Z\s]", "", stripped).lower().strip()
    key = re.sub(r"\s+", " ", key)
    if key in _HEADING_LOOKUP:
        return _HEADING_LOOKUP[key]
    if stripped.isupper() and key in _HEADING_LOOKUP:
        return _HEADING_LOOKUP[key]
    return None


def _looks_like_name(value: str) -> bool:
    if not value or "\n" in value or "\r" in value or _NAME_STOP.search(value) or EMAIL_RE.search(value) or URL_RE.search(value):
        return False
    if any(ch.isdigit() for ch in value):
        return False
    words = [w for w in re.split(r"\s+", value) if w]
    if not (2 <= len(words) <= 4):
        return False
    if len(value) > 60:
        return False
    return all(re.match(r"^[A-Z][A-Za-z'’.\-]+$", w) for w in words)


def _usable_keyword(phrase: str) -> bool:
    tokens = [t.lower() for t in re.findall(r"[A-Za-z][A-Za-z0-9+#.]*", phrase)]
    if not tokens or len(tokens) > 4:
        return False
    if all(t in _STOPWORDS or len(t) < 3 for t in tokens):
        return False
    return True


def _title_skill(item: str) -> str:
    if item.isupper() and len(item) <= 5:
        return item
    if any(c.islower() for c in item) and any(c.isupper() for c in item) and " " not in item:
        return item
    return item.strip()


def _guess_institution(block: str) -> Optional[str]:
    for line in block.split("\n"):
        if find_degree(line):
            continue
        cleaned = line.strip(" -,")
        if cleaned and not re.search(r"(?:19|20)\d{2}", cleaned):
            return cleaned
    return None


def _guess_organization(headline: str) -> Optional[str]:
    for sep in (" at ", " | ", " - ", " – ", "@"):
        if sep in headline:
            parts = [p.strip() for p in re.split(re.escape(sep), headline, maxsplit=1)]
            if len(parts) == 2 and parts[1]:
                return parts[1]
    return None


def _split_blocks(text: str) -> list[str]:
    chunks = re.split(r"\n\s*\n", text.strip())
    blocks = [chunk.strip() for chunk in chunks if chunk.strip()]
    if len(blocks) == 1:
        lines = [line for line in text.split("\n") if line.strip()]
        grouped: list[list[str]] = []
        current: list[str] = []
        for line in lines:
            is_new = bool(re.match(r"^[A-Z].{3,}", line)) and not line.startswith("-")
            if is_new and current and not current[-1].startswith("-"):
                grouped.append(current)
                current = [line]
            else:
                current.append(line)
        if current:
            grouped.append(current)
        if len(grouped) > 1:
            return ["\n".join(g) for g in grouped]
    return blocks


def _bullet_or_line_items(text: str) -> list[str]:
    if not text:
        return []
    items: list[str] = []
    for line in text.split("\n"):
        item = line.strip().lstrip("-").strip()
        if item:
            items.append(item)
    return items


def _first_match(pattern: re.Pattern[str], text: str) -> Optional[str]:
    match = pattern.search(text)
    return match.group(0) if match else None


def _normalize_phone(value: Optional[str]) -> Optional[str]:
    if not value:
        return None
    digits = re.sub(r"\D", "", value)
    if len(digits) < 10 or len(digits) > 15:
        return None
    return value.strip()


def _unique(items: list[str]) -> list[str]:
    seen: set[str] = set()
    out: list[str] = []
    for item in items:
        key = item.lower().strip()
        if key and key not in seen:
            seen.add(key)
            out.append(item.strip())
    return out
